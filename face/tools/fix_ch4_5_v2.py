"""
补修脚本 (修复之前 fix_ch4_5.py 的 bug):
1. 4.3.1 末尾补 detect_faces 示例代码
2. 4.3.2 末尾补 add_person_with_preprocess 示例代码
3. 4.3.3 末尾补 imread 示例代码
4. 5.1 末尾删除重复的 add_person_with_preprocess 代码 (保留 /api/recognize)
5. 5.2 末尾补 preprocess_pipeline + align_face 示例代码 (注意: 之前只插了 1 个, 缺另一个)
6. 5.3.1 末尾补 extract_features_with_preprocess 示例代码
7. 5.3.2 重建 heading + 文字 + 末尾补 __init__ 加载代码
8. 5.3.3 重建 heading + 文字 + 末尾补 imread 代码
9. 5.4 重建 heading + 5 条 List Paragraph

技术修复:
- 避免 doc.add_paragraph 留在末尾的 bug
- 用"先在文档末尾创建, 再 addprevious 移动" 的方式, 但每次都正确移动
"""
import shutil
import docx
from docx.shared import Pt, RGBColor


SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak3'

# 备份当前状态 (修复前快照)
shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)


def make_code(text, indent_pt=14):
    """创建代码段, 返回段落对象 (但暂不加入文档)"""
    p = docx.oxml.ns.qn  # 触发 import 检查
    p = docx.Document().add_paragraph()  # 临时文档, 只为用样式
    # 实际上我们必须用目标文档, 否则样式不匹配
    # 改: 在目标文档末尾创建, 然后再移动
    p = d.add_paragraph()
    p.paragraph_format.left_indent = docx.shared.Pt(indent_pt)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    return p


def insert_before(target_p, new_p):
    """把 new_p 移动到 target_p 之前"""
    target_p._element.addprevious(new_p._element)
    # 新段落已经 addprevious, 自动从原位置 (末尾) 移除
    return new_p


def find_para(text, style_startswith=None):
    """按文本+样式找段落"""
    for p in d.paragraphs:
        if p.text.strip() == text:
            if style_startswith is None or p.style.name.startswith(style_startswith):
                return p
    return None


# ============================================================
# 1. 4.3.1 末尾补 detect_faces 示例
# ============================================================
h_432 = find_para('4.3.2识别准确率优化', 'Heading')
if h_432:
    detect_code = (
        "def detect_faces(self, gray_image, upsample=1):\n"
        "    return self.detector(gray_image, upsample)"
    )
    new_p = make_code(detect_code)
    insert_before(h_432, new_p)
    print('1. 4.3.1 末尾插入 detect_faces ✓')


# ============================================================
# 2. 4.3.2 末尾补多模板平均示例
# ============================================================
h_433 = find_para('4.3.3中文支持优化', 'Heading')
if h_433:
    multi_code = (
        "def add_person_with_preprocess(self, name, image_paths):\n"
        "    feats = []\n"
        "    for p in image_paths:\n"
        "        img = self.imread(p)\n"
        "        if img is None: continue\n"
        "        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)\n"
        "        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)\n"
        "        faces = self.detect_faces(gray)\n"
        "        if not faces: continue\n"
        "        feats.append(self.extract_features_with_preprocess(rgb, faces[0]))\n"
        "    if not feats: return False\n"
        "    avg = np.mean(feats, axis=0)\n"
        "    avg = avg / np.linalg.norm(avg)\n"
        "    self.known_faces.append(avg)\n"
        "    self.known_names.append(name)\n"
        "    return True"
    )
    new_p = make_code(multi_code)
    insert_before(h_433, new_p)
    print('2. 4.3.2 末尾插入 add_person_with_preprocess ✓')


# ============================================================
# 3. 4.3.3 末尾补 imread 示例
# ============================================================
# 4.3.3 后是 5.1 标题
h_51 = find_para('5.1项目成果', 'Heading')
if h_51:
    imread_code_4 = (
        "def imread(self, path):\n"
        "    data = np.fromfile(path, dtype=np.uint8)\n"
        "    return cv2.imdecode(data, cv2.IMREAD_COLOR)"
    )
    new_p = make_code(imread_code_4)
    insert_before(h_51, new_p)
    print('3. 4.3.3 末尾插入 imread ✓')


# ============================================================
# 4. 5.1 末尾删除重复的 add_person_with_preprocess 代码段
# ============================================================
# 5.1 内有两段代码: /api/recognize [227-253] + add_person_with_preprocess [254-272]
# 5.2 heading 之前最后一段是 add_person 重复代码, 要删
h_52 = find_para('5.2技术亮点', 'Heading')
if h_52:
    # 找到 h_52 前面的 Normal (Web) 代码段
    to_remove = []
    for p in d.paragraphs:
        if p.text.strip().startswith('def add_person_with_preprocess'):
            to_remove.append(p)
    # 只删 5.1 区域 (h_52 之前的最后一段)
    if to_remove:
        last = to_remove[-1]
        # 从 d.paragraphs 找 last 的索引
        idx = None
        for i, p in enumerate(d.paragraphs):
            if p._element is last._element:
                idx = i
                break
        if idx is None:
            print('4. ✗ 找不到 last 段索引')
        else:
            # 向上找直到 heading
            block = [last]
            for i in range(idx - 1, -1, -1):
                p = d.paragraphs[i]
                if p.style.name == 'Heading 1' or p.style.name == 'Heading 2':
                    break
                block.append(p)
            for p in block:
                p._element.getparent().remove(p._element)
            print(f'4. 5.1 末尾删除 add_person 重复代码 ({len(block)} 段) ✓')


# ============================================================
# 5. 5.2 末尾补 preprocess_pipeline + align_face
# ============================================================
h_53 = find_para('5.3项目难点与解决方案', 'Heading')
if h_53:
    pre_code = (
        "def preprocess_pipeline(self, image):\n"
        "    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
        "    gray = cv2.equalizeHist(gray)\n"
        "    gray = cv2.GaussianBlur(gray, (5, 5), 0)\n"
        "    return gray"
    )
    new_p1 = make_code(pre_code)
    insert_before(h_53, new_p1)
    align_code = (
        "def align_face(self, rgb_image, face_rect):\n"
        "    shape = self.predictor(rgb_image, face_rect)\n"
        "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
        "    return chip"
    )
    new_p2 = make_code(align_code)
    insert_before(h_53, new_p2)
    print('5. 5.2 末尾插入 preprocess_pipeline + align_face ✓')


# ============================================================
# 6. 5.3.1 末尾补 extract_features_with_preprocess 示例
# ============================================================
# 5.3.1 heading 存在, 文字段后需要补代码
# 找 5.3.1 heading 后面的 "难点" 段, 再后面就是 5.3.2 文字段 (heading 缺失)
# 简化: 找含 "实时性能不足" 的段 (5.3.2 的文字), 在它前面插
text_532 = None
for p in d.paragraphs:
    if '实时性能不足' in p.text:
        text_532 = p
        break
if text_532:
    acc_code = (
        "def extract_features_with_preprocess(self, rgb_image, face_rect):\n"
        "    shape = self.predictor(rgb_image, face_rect)\n"
        "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
        "    descriptor = self.face_model.compute_face_descriptor(\n"
        "        chip, shape, 10)\n"
        "    vec = np.array(descriptor)\n"
        "    vec = vec / np.linalg.norm(vec)\n"
        "    return vec"
    )
    new_p = make_code(acc_code)
    insert_before(text_532, new_p)
    print('6. 5.3.1 末尾插入 extract_features_with_preprocess ✓')


# ============================================================
# 7. 5.3.2 重建 heading + 文字后补 __init__ 代码
# ============================================================
# 找 5.3.3 文字段, 它前面是 5.3.2 文字段
text_533 = None
for p in d.paragraphs:
    if 'Windows系统下OpenCV无法正确处理中文路径' in p.text:
        text_533 = p
        break
if text_533:
    # 在 5.3.3 文字段前面重建 5.3.2 heading + 5.3.2 代码 (但 5.3.2 文字已存在)
    # 实际上文字段已经存在, 我们要做的是: 在 5.3.2 文字段后 (5.3.3 文字段前) 插入代码
    # 然后把缺失的 5.3.2 heading 加到 5.3.2 文字段之前
    init_code = (
        "def __init__(self):\n"
        "    self.detector = dlib.get_frontal_face_detector()\n"
        "    self.predictor = dlib.shape_predictor(\n"
        "        'shape_predictor_68_face_landmarks.dat')\n"
        "    self.face_model = dlib.face_recognition_model_v1(\n"
        "        'dlib_face_recognition_resnet_model_v1.dat')"
    )
    new_p = make_code(init_code)
    insert_before(text_533, new_p)
    # 重建 5.3.2 heading
    text_532_v2 = None
    for p in d.paragraphs:
        if '实时性能不足' in p.text:
            text_532_v2 = p
            break
    if text_532_v2:
        h_p = d.add_paragraph()
        h_p.style = d.styles['Heading 3']
        h_p.add_run('5.3.2性能优化问题')
        insert_before(text_532_v2, h_p)
        print('7. 5.3.2 重建 heading + 末尾插入 __init__ ✓')


# ============================================================
# 8. 5.3.3 重建 heading + 末尾补 imread 代码
# ============================================================
if text_533:
    # 重建 5.3.3 heading
    h_p = d.add_paragraph()
    h_p.style = d.styles['Heading 3']
    h_p.add_run('5.3.3中文路径问题')
    insert_before(text_533, h_p)
    # 末尾补 imread 代码 (在 5.4 段前面)
    # 5.4 文字 "活体检测" 是第一段
    text_54 = None
    for p in d.paragraphs:
        if p.text.strip().startswith('活体检测'):
            text_54 = p
            break
    if text_54:
        imread_code_5 = (
            "def imread(self, path):\n"
            "    data = np.fromfile(path, dtype=np.uint8)\n"
            "    return cv2.imdecode(data, cv2.IMREAD_COLOR)"
        )
        new_p = make_code(imread_code_5)
        insert_before(text_54, new_p)
        print('8. 5.3.3 重建 heading + 末尾插入 imread ✓')


# ============================================================
# 9. 5.4 重建 heading
# ============================================================
text_54 = None
for p in d.paragraphs:
    if p.text.strip().startswith('活体检测'):
        text_54 = p
        break
if text_54:
    h_p = d.add_paragraph()
    h_p.style = d.styles['Heading 3']
    h_p.add_run('5.4未来展望')
    insert_before(text_54, h_p)
    print('9. 5.4 重建 heading ✓')


# ============================================================
# 保存
# ============================================================
d.save(SRC)
print(f'\nSaved: {SRC}')

import os
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')