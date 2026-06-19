"""
修复 4-5 章:
1. 4.3.1 性能优化 - 在 List Paragraph 后补 detect_faces 示例代码
2. 4.3.2 准确率优化 - 在 List Paragraph 后补多模板平均示例代码
3. 4.3.3 中文支持 - 在文字后补 imread 代码 (并删 4.3.1 内的注释 - 没有, 这是其他位置)
4. 5.1 项目成果 - 把 /api/recognize 代码块移到 5 条 List Paragraph 之后
5. 5.2 技术亮点 - 在 5 条 List Paragraph 后补直方图均衡化 + 对齐示例
6. 5.3.1 识别准确率 - 在文字后补示例代码
7. 5.3.2 性能优化 - 在文字后补 __init__ 加载代码
8. 5.3.3 中文路径 - 在文字后补 imread 代码
9. 全文 - 删除 # 开头的注释 (报告里所有代码段的行内注释, 主要是 [231][237][263][274])

注意: 5.1 标题 [223] 后紧跟 4.2.1 add_person_with_preprocess 重复代码 (para[259-277])
      这个其实应该删, 因为 4.2.1 已经有同样的代码. 我也合并清理.
"""
import shutil
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak2'

# 备份
shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)


def make_code_paragraph(doc, text, indent_pt=14):
    """创建一个代码段落, 等宽字体 + 缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = docx.shared.Pt(indent_pt)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    return p


# ============================================================
# Step 1: 删除报告里所有代码中的 # 注释行
# ============================================================
removed_comments = 0
paras_to_remove = []
for i, p in enumerate(d.paragraphs):
    t = p.text
    # 只处理 Normal (Web) 类的代码段
    if p.style.name != 'Normal (Web)':
        continue
    stripped = t.strip()
    if not stripped:
        continue
    # 整行就是 # 注释
    if stripped.startswith('#'):
        paras_to_remove.append(i)
        removed_comments += 1
        continue
    # 行内带 # 注释 (如 'img = self.imread(img_path)        # 支持中文路径')
    if '  #' in t or '\t#' in t or t.lstrip().startswith('return') and '#' in t:
        # 检查行末注释
        # 简单规则: 含 # 且 # 前面是空白/tab
        if '  #' in t or '\t#' in t:
            # 找到第一个 # (不在字符串内, 简化: 取第一个 '  #' 或 '\t#')
            idx = -1
            for marker in ['  #', '\t#']:
                pos = t.find(marker)
                if pos > 0:
                    idx = pos
                    break
            if idx > 0:
                new_text = t[:idx].rstrip()
                # 清空原段落再重写
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                r = p.add_run(new_text)
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(50, 50, 50)
                p.paragraph_format.left_indent = docx.shared.Pt(14)
                removed_comments += 1
                print(f'  [修] para[{i}] 删行内注释: {t[:60]}...')

# 删除整行 # 注释 (从后往前删, 索引不乱)
for idx in reversed(paras_to_remove):
    p = d.paragraphs[idx]
    p._element.getparent().remove(p._element)
print(f'Step 1: 删除 {removed_comments} 处注释')


# ============================================================
# Step 2: 修复 4.3.1 性能优化 - 在 3 条 List Paragraph 后补代码
# ============================================================
# 4.3.1 标题在 para[211], List Paragraph 在 [212][213][214]
# 5.3.2 标题在 para[215] (即 List Paragraph 4 之前)
# 我们要在 para[214] (List Paragraph 3) 之后, 4.3.2 标题之前插入代码

# 重新加载段落列表 (删除后索引变了, 用新索引)
paras = d.paragraphs
idx_432_heading = None
for i, p in enumerate(paras):
    if p.text.strip() == '4.3.2识别准确率优化' and p.style.name.startswith('Heading'):
        idx_432_heading = i
        break

print(f'Step 2: 4.3.2 heading 位置: para[{idx_432_heading}]')

# 在 4.3.2 heading 之前插入 detect_faces 示例代码
detect_code = (
    "def detect_faces(self, gray_image, upsample=1):\n"
    "    return self.detector(gray_image, upsample)"
)
# 由于 docx 难以在中间插入 (只能尾部 add), 我用另一种方法:
# 在 4.3.1 标题 (para[211]) 之后, 把示例代码作为 4.3.1 的最后内容
# 实际是: 找到 4.3.2 heading, 在它前面插入段落 (用 XML 操作)

# 4.3.2 heading 元素
heading_432 = paras[idx_432_heading]._element
# 在 heading 前插入新段落
# 用 docx 创建一个代码段落, 移动到正确位置
new_p = make_code_paragraph(d, detect_code)
new_p._element.addprevious(heading_432)  # 移到 heading 之前
print(f'Step 2: 4.3.1 末尾插入 detect_faces 示例代码')


# ============================================================
# Step 3: 修复 4.3.2 准确率优化 - 在 4 条 List Paragraph 后补多模板平均代码
# ============================================================
# 4.3.3 标题在原 [220] 现在要找
paras = d.paragraphs
idx_433_heading = None
for i, p in enumerate(paras):
    if p.text.strip() == '4.3.3中文支持优化' and p.style.name.startswith('Heading'):
        idx_433_heading = i
        break

print(f'Step 3: 4.3.3 heading 位置: para[{idx_433_heading}]')

multitemplate_code = (
    "def add_person_with_preprocess(self, name, image_paths):\n"
    "    feats = []\n"
    "    for p in image_paths:\n"
    "        img = self.imread(p)\n"
    "        if img is None: continue\n"
    "        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)\n"
    "        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)\n"
    "        faces = self.detect_faces(gray)\n"
    "        if not faces: continue\n"
    "        feats.append(self.extract_features_with_preprocess(rgb, faces[0]))\n"
    "    if not feats: return False\n"
    "    avg = np.mean(feats, axis=0)\n"
    "    avg = avg / np.linalg.norm(avg)\n"
    "    self.known_faces.append(avg)\n"
    "    self.known_names.append(name)\n"
    "    return True"
)
heading_433 = paras[idx_433_heading]._element
new_p = make_code_paragraph(d, multitemplate_code)
new_p._element.addprevious(heading_433)
print(f'Step 3: 4.3.2 末尾插入 add_person_with_preprocess 示例代码')


# ============================================================
# Step 4: 修复 4.3.3 中文支持 - 在文字后补 imread 代码
# ============================================================
paras = d.paragraphs
idx_5_heading = None
for i, p in enumerate(paras):
    if p.text.strip() == '第5章 实训总结' and p.style.name.startswith('Heading'):
        idx_5_heading = i
        break

print(f'Step 4: 第5章 heading 位置: para[{idx_5_heading}]')

imread_code = (
    "def imread(self, path):\n"
    "    data = np.fromfile(path, dtype=np.uint8)\n"
    "    return cv2.imdecode(data, cv2.IMREAD_COLOR)"
)
heading_5 = paras[idx_5_heading]._element
new_p = make_code_paragraph(d, imread_code)
new_p._element.addprevious(heading_5)
print(f'Step 4: 4.3.3 末尾插入 imread 示例代码')


# ============================================================
# Step 5: 修复 5.1 - 把 /api/recognize 代码块从 5.1 标题后挪到 5 条 List Paragraph 之后
# ============================================================
# 当前状态: 5.1 标题 -> /api/recognize 代码 -> 文字 -> 5 条 List Paragraph -> add_person 代码
# 应该是: 5.1 标题 -> 文字 -> 5 条 List Paragraph -> /api/recognize 代码 -> add_person 代码 (2 个)

# 先收集 5.1 内的所有代码段
paras = d.paragraphs
idx_5_1_heading = None
idx_5_2_heading = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == '5.1项目成果' and p.style.name.startswith('Heading'):
        idx_5_1_heading = i
    if t == '5.2技术亮点' and p.style.name.startswith('Heading'):
        idx_5_2_heading = i
        break

print(f'Step 5: 5.1 标题: para[{idx_5_1_heading}], 5.2 标题: para[{idx_5_2_heading}]')

# 收集 5.1 内所有 Normal (Web) 代码段
code_paragraphs_5_1 = []
for i in range(idx_5_1_heading, idx_5_2_heading):
    p = paras[i]
    if p.style.name == 'Normal (Web)':
        code_paragraphs_5_1.append(p)

print(f'  5.1 内有 {len(code_paragraphs_5_1)} 个代码段')

# 把代码段移到 5.1 末尾 (即 5.2 heading 之前)
# 思路: 把所有代码段的 XML 移到 5.2 heading 之前
# 但 python-docx 不会复制, 我们用 element 移动 (addprevious)
heading_52 = paras[idx_5_2_heading]._element
for cp in code_paragraphs_5_1:
    # 1. 从原位置移除
    cp._element.getparent().remove(cp._element)
    # 2. 插入到 5.2 heading 之前
    heading_52.addprevious(cp._element)

print(f'  移动 {len(code_paragraphs_5_1)} 个代码段到 5.1 末尾')


# ============================================================
# Step 6: 5.2 技术亮点后补代码 (2 个示例: 直方图均衡化 + 对齐)
# ============================================================
paras = d.paragraphs
idx_5_2_heading = None
idx_5_3_heading = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == '5.2技术亮点' and p.style.name.startswith('Heading'):
        idx_5_2_heading = i
    if t == '5.3项目难点与解决方案' and p.style.name.startswith('Heading'):
        idx_5_3_heading = i
        break

print(f'Step 6: 5.2 heading: para[{idx_5_2_heading}], 5.3 heading: para[{idx_5_3_heading}]')

# 5.2 内 5 条 List Paragraph (无代码), 在末尾 (5.3 heading 前) 插入 2 个示例
preprocess_code = (
    "def preprocess_pipeline(self, image):\n"
    "    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
    "    gray = cv2.equalizeHist(gray)\n"
    "    gray = cv2.GaussianBlur(gray, (5, 5), 0)\n"
    "    return gray"
)
align_code = (
    "def align_face(self, rgb_image, face_rect):\n"
    "    shape = self.predictor(rgb_image, face_rect)\n"
    "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
    "    return chip"
)
heading_53 = paras[idx_5_3_heading]._element
new_p1 = make_code_paragraph(d, preprocess_code)
new_p1._element.addprevious(heading_53)
new_p2 = make_code_paragraph(d, align_code)
new_p2._element.addprevious(heading_53)
print(f'Step 6: 5.2 末尾插入 preprocess_pipeline + align_face 示例代码')


# ============================================================
# Step 7: 5.3.1 识别准确率 - 在文字后补示例代码
# ============================================================
paras = d.paragraphs
idx_5_3_1_heading = None
idx_5_3_2_heading = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == '5.3.1识别准确率问题' and p.style.name.startswith('Heading'):
        idx_5_3_1_heading = i
    if t == '5.3.2性能优化问题' and p.style.name.startswith('Heading'):
        idx_5_3_2_heading = i
        break

print(f'Step 7: 5.3.1 heading: para[{idx_5_3_1_heading}], 5.3.2 heading: para[{idx_5_3_2_heading}]')

accuracy_code = (
    "def extract_features_with_preprocess(self, rgb_image, face_rect):\n"
    "    shape = self.predictor(rgb_image, face_rect)\n"
    "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
    "    descriptor = self.face_model.compute_face_descriptor(\n"
    "        chip, shape, 10)\n"
    "    vec = np.array(descriptor)\n"
    "    vec = vec / np.linalg.norm(vec)\n"
    "    return vec"
)
heading_532 = paras[idx_5_3_2_heading]._element
new_p = make_code_paragraph(d, accuracy_code)
new_p._element.addprevious(heading_532)
print(f'Step 7: 5.3.1 末尾插入 extract_features_with_preprocess 示例代码')


# ============================================================
# Step 8: 5.3.2 性能优化 - 在文字后补 __init__ 加载代码
# ============================================================
paras = d.paragraphs
idx_5_3_2_heading = None
idx_5_3_3_heading = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == '5.3.2性能优化问题' and p.style.name.startswith('Heading'):
        idx_5_3_2_heading = i
    if t == '5.3.3中文路径问题' and p.style.name.startswith('Heading'):
        idx_5_3_3_heading = i
        break

print(f'Step 8: 5.3.2 heading: para[{idx_5_3_2_heading}], 5.3.3 heading: para[{idx_5_3_3_heading}]')

init_code = (
    "def __init__(self):\n"
    "    self.detector = dlib.get_frontal_face_detector()\n"
    "    self.predictor = dlib.shape_predictor(\n"
    "        'shape_predictor_68_face_landmarks.dat')\n"
    "    self.face_model = dlib.face_recognition_model_v1(\n"
    "        'dlib_face_recognition_resnet_model_v1.dat')"
)
heading_533 = paras[idx_5_3_3_heading]._element
new_p = make_code_paragraph(d, init_code)
new_p._element.addprevious(heading_533)
print(f'Step 8: 5.3.2 末尾插入 __init__ 加载模型示例代码')


# ============================================================
# Step 9: 5.3.3 中文路径 - 在文字后补 imread 代码 (与 4.3.3 相同, 但放在 5.3.3)
# ============================================================
paras = d.paragraphs
idx_5_3_3_heading = None
idx_5_4_heading = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t == '5.3.3中文路径问题' and p.style.name.startswith('Heading'):
        idx_5_3_3_heading = i
    if t == '5.4未来展望' and p.style.name.startswith('Heading'):
        idx_5_4_heading = i
        break

print(f'Step 9: 5.3.3 heading: para[{idx_5_3_3_heading}], 5.4 heading: para[{idx_5_4_heading}]')

chinese_code = (
    "def imread(self, path):\n"
    "    data = np.fromfile(path, dtype=np.uint8)\n"
    "    return cv2.imdecode(data, cv2.IMREAD_COLOR)"
)
heading_54 = paras[idx_5_4_heading]._element
new_p = make_code_paragraph(d, chinese_code)
new_p._element.addprevious(heading_54)
print(f'Step 9: 5.3.3 末尾插入 imread 示例代码')


# ============================================================
# 保存
# ============================================================
d.save(SRC)
print(f'\nSaved: {SRC}')

import os
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')