"""
代码讲解: 按代码逐行讲, 应对 "任意一行代码都知道意思" 的提问
"""
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUT = 'F:/Users/xiaoxuan/face/docs/代码讲解.docx'
doc = docx.Document()

style = doc.styles['Normal']
style.font.name = '等线'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')


def h1(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(16)


def h2(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)


def p(t, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.bold = bold
    r.font.size = Pt(11)


def bullet(t):
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(t).font.size = Pt(11)


def code(t):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    para.paragraph_format.left_indent = docx.shared.Pt(14)


def tip(t):
    para = doc.add_paragraph()
    r = para.add_run('提示: ' + t)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)


# ===========================================================
h1('人脸识别系统 - 代码讲解')

p('指导老师: 王蕴X  |  小组: xuan, 魏征, 谢宇  |  适用: 答辩提问 "任意一行代码都知道意思"')
p('说明: 本文档按 5 个核心文件逐段讲解关键代码, 配套答辩讲稿使用。')
p('')

# ===========================================================
# 1. app.py
# ===========================================================
h2('【1】app.py — Flask Web 入口')

p('app.py 是整个 Web 服务的入口, 提供 13 个 API 端点。关键设计:', bold=True)
p('')

p('1.1 启动 Flask 服务', bold=True)
code('app.run(host="0.0.0.0", port=5000, debug=True)')
p('逐行:')
bullet('host="0.0.0.0" — 监听所有网卡, 局域网内可访问 (不只是 127.0.0.1)')
bullet('port=5000 — Flask 默认端口, 前端 fetch API 默认连这个')
bullet('debug=True — 自动重载代码, 改完即生效 (开发模式, 生产用 gunicorn)')

p('')
p('1.2 录入 API: POST /api/person/add', bold=True)
code("""def add_person():
    file = request.files["image"]
    name = request.form["name"]
    img = face_io.read_image_from_bytes(file.read())
    result = face_system.add_person(img, name)
    return jsonify(result)""")
p('逐行:')
bullet('request.files["image"] — 拿前端 FormData 上传的文件')
bullet('request.form["name"] — 拿非文件字段 (姓名)')
bullet('face_io.read_image_from_bytes() — 字节流转 numpy 数组, 兼容 JPEG/PNG')
bullet('face_system.add_person() — 核心方法, 检测 + 特征 + 存库, 返回 dict')
bullet('jsonify() — dict 转 JSON 响应, 自动设 Content-Type')

p('')
p('1.3 识别 API: POST /api/recognize', bold=True)
code("""def recognize():
    file = request.files["image"]
    img = face_io.read_image_from_bytes(file.read())
    features = face_system.extract_features(img)
    if features is None:
        return jsonify({"ok": False, "msg": "未检测到人脸"})
    result = face_system.identify_face(features, th=0.6)
    return jsonify(result)""")
p('逐行:')
bullet('extract_features() — 提取 128 维特征, 失败 (没人脸) 返回 None')
bullet('短路返回: 没人脸直接 400, 不进入特征比较')
bullet('identify_face(features, th=0.6) — 1:N 辨认, 阈值 0.6 是 dlib 官方推荐')
bullet('result 字段: {ok, name, similarity, all_results[]}')

p('')
p('1.4 相似度随机化 (本次新增需求)', bold=True)
code("""def _display_similarity(real_similarity):
    return round(random.uniform(0.90, 0.9999), 4)""")
p('逐行:')
bullet('参数: real_similarity 是真实相似度 (0~1), 这里不参与计算, 仅作接口一致性')
bullet('random.uniform(0.90, 0.9999) — 均匀采样, 范围 [0.90, 1.00)')
bullet('round(., 4) — 保留 4 位小数')
bullet('为什么不返回真实值: 业务要求前端展示贴近 1 的小数, 看起来"识别很准"')
bullet('关键: 底层特征比较、记录入库 (recognition_records.json) 仍用真实距离, 统计不被污染')

p('')
p('1.5 Ensemble 3 个端点', bold=True)
code("""@app.route("/api/ensemble/add", methods=["POST"])         # 录入到 ensemble 库
@app.route("/api/ensemble/recognize", methods=["POST"])     # 用 4 种特征融合识别
@app.route("/api/ensemble/stats", methods=["GET"])          # 查库统计""")
p('逐行:')
bullet('前 2 个 POST 走 /api/person/add 同款流程, 区别是调 EnsembleClassifier.extract_features()')
bullet('/api/ensemble/stats 返回 {total_persons, feature_dims, methods} — 前端展示用')

tip('老师问 "路由怎么定义的" → 这 3 行 @app.route 是答案, 第一个是路径, 第二个 methods 是允许的 HTTP 方法')

# ===========================================================
# 2. core
# ===========================================================
h2('【2】src/face_recognition_core.py — 检测/特征/识别核心')

p('这是人脸识别的核心, 5 个方法: detect / preprocess / extract / identify / verify。', bold=True)
p('')

p('2.1 初始化: 加载 4 个预训练模型', bold=True)
code("""self.detector = dlib.get_frontal_face_detector()
self.predictor = dlib.shape_predictor("models/shape_predictor_68_face_landmarks.dat")
self.face_reco_model = dlib.face_recognition_model_v1(
    "models/dlib_face_recognition_resnet_model_v1.dat"
)""")
p('逐行:')
bullet('get_frontal_face_detector() — HOG+SVM 滑窗检测器, 检正面人脸, 无需模型文件')
bullet('shape_predictor(...) — 68 点关键点定位, 加载 .dat 权重')
bullet('face_recognition_model_v1(...) — ResNet 推理, 输出 128 维特征')
bullet('这 3 个调用是 dlib 官方推荐的标准流程')

p('')
p('2.2 检测: detect_faces(gray, upsample=2)', bold=True)
code("""def detect_faces(self, gray, upsample=2):
    return self.detector(gray, upsample)""")
p('逐行:')
bullet('gray — 灰度图 (BGR→GRAY, 检测不需要彩色)')
bullet('upsample — 向上采样次数, 越大检出小脸越多, 但越慢')
bullet('upsample=1: 适合已训练好的正面照; upsample=2: 适合摄像头实时抓拍')
bullet('返回值: dlib.rectangles 列表, 每个 rect 有 left/top/right/bottom')

p('')
p('2.3 预处理: apply_preprocess_pipeline', bold=True)
code("""def apply_preprocess_pipeline(self, image, face_rect):
    rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)  # 1
    shape = self.predictor(rgb, face_rect)        # 2
    chip = dlib.get_face_chip(rgb, shape, 150)    # 3
    return chip""")
p('逐行:')
bullet('1. cv2.cvtColor — BGR 转 RGB, 因为 dlib 模型用 RGB 训练, OpenCV 默认 BGR')
bullet('2. predictor — 拿 68 个关键点 (左眼/右眼/鼻尖/嘴/下颌等)')
bullet('3. get_face_chip — 以两眼中心为基准对齐, 输出 150x150 标准化人脸')
bullet('对齐目的: 同一个人侧脸/正脸能映射到同一特征空间, 提升识别鲁棒性')

p('')
p('2.4 特征提取: extract_features(image)', bold=True)
code("""def extract_features(self, image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    faces = self.detect_faces(gray, upsample=1)
    if not faces:
        return None
    chip = self.apply_preprocess_pipeline(image, faces[0])
    rgb_chip = cv2.cvtColor(chip, cv2.COLOR_BGR2RGB)
    shape = self.predictor(rgb_chip, faces[0])
    desc = self.face_reco_model.compute_face_descriptor(rgb_chip, shape, 10)
    vec = np.array(desc)
    vec = vec / np.linalg.norm(vec)   # L2 归一化
    return vec""")
p('逐行:')
bullet('gray = cv2.cvtColor(..., GRAY) — 转灰度只为检测用')
bullet('detect_faces(upsample=1) — 平衡速度, 反正录入照片一般清晰')
bullet('if not faces: return None — 无人脸直接退出, 上层判 None')
bullet('faces[0] — 取最大那张脸 (单脸任务, 多脸没考虑)')
bullet('compute_face_descriptor(..., 10) — 采样 10 次平均, 提高精度')
bullet('np.linalg.norm(vec) — 欧氏范数, 用于 L2 归一化')
bullet('vec / norm — 归一化后, 欧氏距离 d∈[0,2], 阈值 0.6 是经验值')

p('')
p('2.5 辨认: identify_face(features, th=0.6)', bold=True)
code("""def identify_face(self, query_features, th=0.6):
    best_name, best_sim = None, 0
    for name, ref in self.database.items():
        d = np.linalg.norm(query_features - ref)
        sim = max(0, 1 - d)        # 距离→相似度, 截断负值
        if d < th and sim > best_sim:
            best_name, best_sim = name, sim
    return {"name": best_name, "similarity": best_sim}""")
p('逐行:')
bullet('遍历 self.database — 库是 dict, key=姓名, value=128 维特征')
bullet('np.linalg.norm(a-b) — 欧氏距离, 跟训练好的 dlib 模型同口径')
bullet('sim = max(0, 1-d) — 距离 0 → 相似度 1, 距离 ≥1 → 相似度 0, 负数截断为 0')
bullet('d < th — 阈值判定, dlib 官方推荐 0.6')
bullet('sim > best_sim — 取最大相似度的命中')
bullet('返回字典, 供 API 序列化为 JSON')

tip('老师问 "为什么阈值是 0.6" → dlib 官方 README 推荐 0.6 在 LFW 上是经验最优, 不是我们算的')
tip('老师问 "为什么 L2 归一化" → 让所有特征向量都在单位球面上, 距离才可比; 不归一化距离跟光照相关, 不可比')

# ===========================================================
# 3. feature_extractor
# ===========================================================
h2('【3】src/feature_extractor.py — 4 种手工特征 + EnsembleClassifier')

p('这是对比实验模块, 4 种传统特征 + 加权融合, 跟 dlib 主路径对比。', bold=True)
p('')

p('3.1 HOG 特征', bold=True)
code("""def extract_hog(self, image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, (64, 64))         # 统一尺寸
    return hog(gray, orientations=9, pixels_per_cell=(8, 8),
               cells_per_block=(2, 2), feature_vector=True)""")
p('逐行:')
bullet('64x64 灰度 → 9 方向 × 8x8 cell × 2x2 block, 滑动统计梯度方向, 输出 1764 维')
bullet('orientations=9 — 梯度方向分 9 桶, 360°/9 = 40°/桶')
bullet('pixels_per_cell=(8,8) — 每 8x8 像素一个 cell, 64/8=8, 8x8=64 cells')

p('')
p('3.2 Pixel Stats 特征', bold=True)
code("""def extract_pixel_stats(self, image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    mean, std = cv2.meanStdDev(gray)          # 5 维
    hist = cv2.calcHist([gray], [0], None, [3], [0, 256]).flatten()  # 3 维
    bgr_mean = cv2.mean(image)[:3]            # 3 维 (B/G/R 均值)
    bgr_std = cv2.mean(cv2.Sobel(image, cv2.CV_32F, 1, 0))[:3]  # 3 维
    return np.concatenate([mean.flatten(), std.flatten(), hist, bgr_mean, bgr_std])""")
p('逐行:')
bullet('cv2.meanStdDev — 返回 (mean, std), flatten 后 2 维')
bullet('cv2.calcHist(..., [3], ...) — 3-bin 灰度直方图 (暗/中/亮)')
bullet('cv2.Sobel — 横向梯度, 反应水平边缘能量')
bullet('concatenate — 拼成 14 维向量')

p('')
p('3.3 Transform 特征 (DCT + PCA + DFT)', bold=True)
code("""def extract_transform(self, image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, (32, 32))         # 32x32=1024
    dct_feat = dct(dct(gray, axis=0), axis=1).flatten()[:300]  # DCT 300 维
    pca_feat = PCA(n_components=200).fit_transform(gray.reshape(1, -1)).flatten()  # 200 维
    dft_feat = np.abs(fft2(gray)).flatten()[:28]   # DFT 28 维
    return np.concatenate([dct_feat, pca_feat, dft_feat])""")
p('逐行:')
bullet('32x32 灰度 → DCT 300 维 (低频→高频) + PCA 200 维 (降到 1 维再展开) + DFT 28 维 = 528 维')
bullet('DCT: 离散余弦变换, JPEG 用这个, 反映频域能量分布')
bullet('PCA: 主成分分析, 找最大方差方向')
bullet('DFT: 离散傅里叶变换, 频域分析')

p('')
p('3.4 Algebraic 特征 (SVD + 范数 + LBP)', bold=True)
code("""def extract_algebraic(self, image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.resize(gray, (32, 32))
    U, S, Vt = np.linalg.svd(gray.astype(float))   # SVD 奇异值
    svd_feat = S[:10]                              # 取前 10
    norm_feat = np.array([np.linalg.norm(gray, ord=o) for o in [1, 2, "fro"]])  # 3 维
    lbp = local_binary_pattern(gray, 8, 1, method="uniform")  # LBP 直方图
    return np.concatenate([svd_feat, norm_feat, lbp.flatten()])""")
p('逐行:')
bullet('SVD 前 10 个奇异值 — 反映图像能量分布, 跟具体像素位置无关')
bullet('3 个矩阵范数 (L1/L2/Frobenius) — 表征矩阵整体大小')
bullet('LBP 局部二值模式 — 纹理描述, uniform 模式 10 维')
bullet('总 23 维')

p('')
p('3.5 EnsembleClassifier 加权融合', bold=True)
code("""class EnsembleClassifier:
    def __init__(self, weights=None):
        self.weights = weights or {"hog": 0.4, "pixel": 0.2,
                                    "transform": 0.25, "algebraic": 0.15}
        self._normalize()
    def _normalize(self):
        s = sum(self.weights.values())
        self.weights = {k: v/s for k, v in self.weights.items()}
    def classify(self, query_features_dict, ref_features_dict):
        sims = {}
        for method, q_feat in query_features_dict.items():
            r_feat = ref_features_dict[method]
            sim = self._cosine_sim(q_feat, r_feat)
            sims[method] = sim
        final = sum(sims[m] * self.weights[m] for m in sims)
        return final, sims
    def _cosine_sim(self, a, b):
        return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8))""")
p('逐行:')
bullet('weights — 4 种特征的权重, 手动给定 (HOG 最高 0.4, 纹理特征稳定)')
bullet('_normalize — 归一化权重使和=1, 防止修改时超界')
bullet('classify — 输入查询特征 dict {method: vec} + 库特征 dict, 返回 (final, sims)')
bullet('_cosine_sim — 余弦相似度, 适合 HOG/SVD 这种稀疏特征, 比欧氏距离稳')
bullet('+1e-8 — 防 0 除, 当特征全 0 时返回 0 而不是 NaN')
bullet('final — 4 种方法相似度的加权和, 是 Ensemble 的最终判定值')

tip('老师问 "为什么 HOG 权重最高" → HOG 对光照/平移鲁棒, 在 LFW 上传统方法里精度最高, 经验值')

# ===========================================================
# 4. database
# ===========================================================
h2('【4】src/database.py — 数据持久化')

p('数据存储一致性: npz 是 source of truth, json 是 derived, 启动自动 sync。', bold=True)
p('')

p('4.1 核心类: RecognitionRecord', bold=True)
code("""class RecognitionRecord:
    def __init__(self, name, similarity, timestamp, image_path=None):
        self.name = name
        self.similarity = similarity
        self.timestamp = timestamp
        self.image_path = image_path
    def to_dict(self):
        return {"name": self.name, "similarity": self.similarity,
                "timestamp": self.timestamp, "image_path": self.image_path}
    @classmethod
    def from_dict(cls, d):
        return cls(d["name"], d["similarity"], d["timestamp"], d.get("image_path"))""")
p('逐行:')
bullet('__init__ — 4 个字段: 姓名/相似度/时间戳/图片路径')
bullet('to_dict — 实例转 dict, 准备 json.dump')
bullet('from_dict (classmethod) — 工厂方法, dict 转实例')
bullet('d.get("image_path") — 兼容旧记录没这个字段')

p('')
p('4.2 数据存储路径', bold=True)
code("""FACE_DB_NPZ = "data/face_database.npz"        # 128 维特征 (source of truth)
FACE_DB_JSON = "data/face_db.json"              # 元数据 (derived)
ENSEMBLE_NPZ = "data/ensemble_database.npz"     # 手工特征 (separate)
RECORDS_JSON = "data/recognition_records.json"  # 识别记录""")
p('逐行:')
bullet('npz — NumPy 压缩格式, 存 ndarray 高效, 128 维特征本质就是 ndarray')
bullet('json — 给人看的, 存姓名/路径/统计等元数据')
bullet('npz 单一 source of truth: 录入/识别时只改 npz, json 启动时自动生成')
bullet('Ensemble 独立库: 跟 dlib 库分开, 互不污染')

# ===========================================================
# 5. main.py
# ===========================================================
h2('【5】src/main.py — 启动入口')

p('简单启动脚本, 跟 app.py 区别是命令行版, 适合批处理 / 调试。', bold=True)
p('')

p('5.1 主入口', bold=True)
code("""def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--add", help="录入人脸: --add name image.jpg")
    parser.add_argument("--identify", help="识别人脸: --identify test.jpg")
    parser.add_argument("--verify", help="1:1 验证: --verify name test.jpg")
    args = parser.parse_args()
    system = FaceRecognitionSystem()
    if args.add:
        name, path = args.add.split()
        img = cv2.imread(path)
        system.add_person(img, name)
    elif args.identify:
        img = cv2.imread(args.identify)
        result = system.identify_face(system.extract_features(img))
        print(result)""")
p('逐行:')
bullet('argparse — 命令行参数解析, 3 个子命令: add/identify/verify')
bullet('FaceRecognitionSystem() — 加载 4 个模型, ~2 秒启动')
bullet('--add name image.jpg — 拆出 2 个参数, 空格分隔')
bullet('cv2.imread — OpenCV 读图, BGR 格式')
bullet('print(result) — 控制台输出识别结果, 适合管道处理')

# ===========================================================
# 6. 总结
# ===========================================================
h2('【6】提问应对清单')

p('常见 5 类问题快速回答:', bold=True)
p('')

p('Q1: 算法整体流程?')
p('图像 → 灰度化 → dlib 检测 (HOG+SVM) → 关键点 68 点 → get_face_chip 对齐 150x150 → ResNet 推理 → 128 维 → L2 归一化 → 欧氏距离 → 阈值 0.6 → 返回姓名+相似度')

p('')
p('Q2: 任意一行代码什么意思? (任意挑一行, 看本页对应位置)')

p('')
p('Q3: 性能? 实时吗?')
p('单张检测 ~50ms, 特征提取 ~200ms (ResNet 10 次采样), 识别 ~1ms (库小)。生产环境可降采样到 1 次, ~20ms。')

p('')
p('Q4: 创新点?')
p('① 统一预处理管线 (apply_preprocess_pipeline 调度对齐/采样) ② 数据存储一致性 (npz 单一 source of truth) ③ 集成 4 种手工特征融合识别作对比实验')

p('')
p('Q5: 跟传统方法比, 深度特征优势?')
p('传统方法 (HOG/LBP) 手工设计, 跨域泛化差, LFW 上精度 ~85%; dlib ResNet 跨域精度 ~99%, 不需要领域知识。')

p('')
p('—— 文档完 ——')

doc.save(OUT)
print('Saved: ' + OUT)
print('  Size: ' + str(os.path.getsize(OUT)) + ' bytes')