"""
微调: 5.2 末尾示例代码顺序 (preprocess 应在 align_face 之前)
"""
import shutil
import docx
from docx.shared import Pt, RGBColor

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak9'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)

# 找 5.3 heading
h_53 = None
for p in paras:
    if p.text.strip() == '5.3项目难点与解决方案' and p.style.name.startswith('Heading'):
        h_53 = p
        break
print(f'5.3 heading: {h_53 is not None}')

# 找当前 [265] align_face 段和 [266] preprocess 段
align_p = None
pre_p = None
for p in d.paragraphs:
    t = p.text.strip()
    if t.startswith('def align_face'):
        align_p = p
    elif t.startswith('def preprocess_pipeline'):
        pre_p = p

print(f'align_p: {align_p is not None}, pre_p: {pre_p is not None}')

if align_p and pre_p:
    # 顺序错: 现在 [align, pre]
    # 期望: [pre, align]
    # 把 pre 移动到 align 之前
    align_p._element.addprevious(pre_p._element)
    print('调整顺序: preprocess_pipeline 移到 align_face 之前 ✓')

d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')