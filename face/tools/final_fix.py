"""
一次性完整修复 4-5 章:
A. 重建 4.3.2 heading (在"图像预处理"前)
B. 重建 4.3.3 heading (在中文路径文字段前)
C. 4.3.1 末尾补 detect_faces 示例
D. 4.3.2 末尾补 add_person_with_preprocess 示例
E. 4.3.3 末尾补 imread 示例 (已经有了, 跳过)
F. 5.1 末尾保留 /api/recognize, 删重复 add_person (已删)
G. 5.2 末尾补 preprocess + align_face 示例
"""
import shutil
import docx
from docx.shared import Pt, RGBColor

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak7'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)


def make_heading(text, level=3):
    """创建 heading 段并加入文档, 返回段落对象"""
    p = d.add_paragraph()
    p.style = d.styles[f'Heading {level}']
    r = p.add_run(text)
    return p


def make_code(text):
    """创建代码段并加入文档, 返回段落对象"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = docx.shared.Pt(14)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    return p


def insert_before(target_p, new_p):
    target_p._element.addprevious(new_p._element)


def find_para_by_text(text, style_prefix='Heading'):
    for p in d.paragraphs:
        if p.text.strip() == text and p.style.name.startswith(style_prefix):
            return p
    return None


def find_para_contains(substring, style=None):
    for p in d.paragraphs:
        if substring in p.text:
            if style is None or p.style.name == style:
                return p
    return None


# ============================================================
# A. 重建 4.3.2 heading
# ============================================================
target_432 = find_para_contains('图像预处理——通过直方图均衡化')
if target_432:
    h = make_heading('4.3.2识别准确率优化', level=3)
    insert_before(target_432, h)
    print('A. 重建 4.3.2 heading ✓')


# ============================================================
# B. 重建 4.3.3 heading
# ============================================================
target_433 = find_para_contains('OpenCV在Windows系统下无法直接处理中文路径')
if target_433:
    h = make_heading('4.3.3中文支持优化', level=3)
    insert_before(target_433, h)
    print('B. 重建 4.3.3 heading ✓')


# ============================================================
# C. 4.3.1 末尾补 detect_faces
# 找到 4.3.2 heading (刚刚创建), 在它之前插入
# ============================================================
h_432 = find_para_by_text('4.3.2识别准确率优化')
if h_432:
    detect_code = (
        "def detect_faces(self, gray_image, upsample=1):\n"
        "    return self.detector(gray_image, upsample)"
    )
    new_p = make_code(detect_code)
    insert_before(h_432, new_p)
    print('C. 4.3.1 末尾补 detect_faces ✓')


# ============================================================
# D. 4.3.2 末尾补多模板平均
# ============================================================
h_433 = find_para_by_text('4.3.3中文支持优化')
if h_433:
    multi_code = (
        "def add_person_with_preprocess(self, name, image_paths):\n"
        "    feats = []\n"
        "    for p in image_paths:\n"
        "        img = self.imread(p)\n"
        "        if img is None: continue\n"
        "        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)\n"
        "        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)\n"
        "        faces = self.detect_faces(gray)\n"
        "        if not faces: continue\n"
        "        feats.append(self.extract_features_with_preprocess(rgb, faces[0]))\n"
        "    if not feats: return False\n"
        "    avg = np.mean(feats, axis=0)\n"
        "    avg = avg / np.linalg.norm(avg)\n"
        "    self.known_faces.append(avg)\n"
        "    self.known_names.append(name)\n"
        "    return True"
    )
    new_p = make_code(multi_code)
    insert_before(h_433, new_p)
    print('D. 4.3.2 末尾补 add_person_with_preprocess ✓')


# ============================================================
# G. 5.2 末尾补 preprocess + align_face
# ============================================================
h_53 = find_para_by_text('5.3项目难点与解决方案')
if h_53:
    pre_code = (
        "def preprocess_pipeline(self, image):\n"
        "    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
        "    gray = cv2.equalizeHist(gray)\n"
        "    gray = cv2.GaussianBlur(gray, (5, 5), 0)\n"
        "    return gray"
    )
    align_code = (
        "def align_face(self, rgb_image, face_rect):\n"
        "    shape = self.predictor(rgb_image, face_rect)\n"
        "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
        "    return chip"
    )
    p_pre = make_code(pre_code)
    p_align = make_code(align_code)
    # 先移 p_align 到 h_53 前 (此时 h_53 前只有 p_align)
    insert_before(h_53, p_align)
    # 再移 p_pre 到 h_53 前 (此时 h_53 前是 p_pre, p_align)
    insert_before(h_53, p_pre)
    print('G. 5.2 末尾补 preprocess + align_face ✓')


d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')