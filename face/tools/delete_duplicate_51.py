"""
直接删 5.1 末尾 add_person 重复代码:
- 5.1 内 5 条 List Paragraph 在前
- /api/recognize 代码在中间 (要保留)
- add_person 重复代码在最后 (要删)
- 5.2 heading 在末尾

策略: 从 5.2 heading 之前往前找, 删到 'return True' (add_person 末尾) 为止
"""
import shutil
import docx

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak4'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)

# 找 5.2 heading 位置
idx_52 = None
for i, p in enumerate(paras):
    if p.text.strip() == '5.2技术亮点' and p.style.name.startswith('Heading'):
        idx_52 = i
        break
print(f'5.2 heading at: {idx_52}')

# 从 5.2 heading 往前找 add_person 起始段
idx_add_person_start = None
for i in range(idx_52 - 1, -1, -1):
    t = paras[i].text.strip()
    if t.startswith('def add_person_with_preprocess'):
        idx_add_person_start = i
        break
print(f'add_person start at: {idx_add_person_start}')

# 删 idx_add_person_start 到 idx_52-1 之间的所有段 (含)
if idx_add_person_start is not None and idx_52 is not None:
    to_delete = paras[idx_add_person_start:idx_52]
    for p in to_delete:
        p._element.getparent().remove(p._element)
    print(f'Deleted {len(to_delete)} paragraphs (add_person 重复代码段)')

d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')