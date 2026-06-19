"""
修改实训报告 docx:
1. 删除 GUI 相关描述
2. 把代码片段替换为项目实际代码
3. 把接口表改为 Flask 路由
4. 把技术栈、数据库分析、界面设计等改为 Web 端
"""
import docx
from docx.oxml.ns import qn
import copy
import shutil
import os

SRC = 'F:/Users/xiaoxuan/face/实训报告模板.docx'
DST = 'F:/Users/xiaoxuan/face/实训报告_修改版.docx'

# 复制源文件
shutil.copy(SRC, DST)

doc = docx.Document(DST)

# 工具函数: 找到包含某文本的段落索引
def find_paragraph_index(text_match, exact=False):
    for i, p in enumerate(doc.paragraphs):
        if exact:
            if p.text.strip() == text_match:
                return i
        else:
            if text_match in p.text:
                return i
    return -1

# 工具函数: 替换段落的文本(保留样式)
def replace_paragraph_text(p, new_text):
    # 保存原样式
    style = p.style
    # 清除所有 run
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # 加一个新 run
    p.add_run(new_text)
    return p

# 工具函数: 删除段落(从 XML 中移除)
def remove_paragraph(p):
    p._element.getparent().remove(p._element)

# 工具函数: 找到从 start_idx 开始的连续 N 个段落,返回索引列表
def slice_paragraphs(start_idx, end_idx):
    # 在当前文档状态下从 start_idx 找,找到文本为某个关键词的
    return list(range(start_idx, end_idx))

# ============================================================
# 1) 修改 4.1.2 技术栈
# ============================================================
# P135: "GUI框架：Tkinter" -> "Web框架：Flask"
idx = find_paragraph_index('GUI框架')
if idx >= 0:
    replace_paragraph_text(doc.paragraphs[idx], 'Web框架：Flask')

# 1.2 节里说"用户可以通过图形化界面操作" -> Web 化
idx = find_paragraph_index('用户可以通过图形化界面进行人脸录入')
if idx >= 0:
    replace_paragraph_text(doc.paragraphs[idx],
        '用户可以通过浏览器访问 Web 界面进行人脸录入、识别、人员管理等操作。')

# ============================================================
# 2) 修改 3.3 数据库分析
# ============================================================
idx = find_paragraph_index('系统采用文件系统存储人脸数据')
if idx >= 0:
    replace_paragraph_text(doc.paragraphs[idx],
        '系统采用 NumPy 二进制文件（face_database.npz）存储 128 维人脸特征向量，'
        '并以 JSON 文件（face_db.json）维护人员姓名、特征元数据、图像路径和录入时间。'
        '识别记录单独存储在 recognition_records.json 中，包含每次识别的时间、姓名、'
        '置信度。数据库加载与持久化通过 FaceRecognitionCore.load_database() 与 '
        'save_database() 接口完成，JSON 数据库通过 FaceDatabase 与 RecognitionRecord '
        '两个类管理。')

# ============================================================
# 3) 修改 3.4 接口设计表 (T2)
# ============================================================
# 接口表在 T2
target_table = doc.tables[2]
# 把整张表重写
# 7 行(原表) + 1 表头 = 8 行
# 先清掉所有行(保留表头)
for row in list(target_table.rows)[1:]:
    row._element.getparent().remove(row._element)

new_rows = [
    ('接口名称', '方法', '功能描述'),
    ('/api/person/add', 'POST', '上传图片并录入新人脸特征'),
    ('/api/person/delete', 'POST', '删除已录入的人员及其特征'),
    ('/api/persons', 'GET', '获取已录入人员列表'),
    ('/api/recognize', 'POST', '识别上传图片中的人脸（支持 1:1 确认 / 1:N 辨认）'),
    ('/api/records', 'GET', '获取最近的识别记录'),
    ('/api/records/clear', 'POST', '清空所有识别记录'),
    ('/api/features', 'GET', '获取当前支持的特征提取方式信息'),
    ('/api/features/extract', 'POST', '提取单张图片的特征向量用于实验分析'),
    ('/api/stats', 'GET', '获取人员数量、记录数量等统计信息'),
    ('/api/persons/refresh', 'POST', '重新加载人脸数据库'),
]

# 添新行
for cells in new_rows[1:]:
    row = target_table.add_row()
    for i, txt in enumerate(cells):
        row.cells[i].text = txt

# ============================================================
# 4) 修改 3.5 用户界面设计
# ============================================================
# 找到 3.5 节起始段
idx_p119 = find_paragraph_index('系统提供以下主要功能接口')
# 3.5 节是从 "系统采用图形化用户界面" 开始的
idx_start = find_paragraph_index('系统采用图形化用户界面')
# 把 3.5 整段(从 P120 到 P125)替换
if idx_start >= 0:
    # 删除原 P120~P125
    # 实际索引: idx_start 到 idx_start+5
    for _ in range(6):
        p_to_del = doc.paragraphs[idx_start]
        remove_paragraph(p_to_del)
    # 重新插入新内容
    new_paras = [
        '系统采用浏览器端 Web 界面（基于 HTML5 + CSS3 + 原生 JavaScript，'
        '通过 Fetch API 与 Flask 后端通信），主要功能区域包括：',
        '1. 主页 / 控制台：上传待识别图片、选择识别模式（1:1 确认 / 1:N 辨认）、'
        '选择目标人员（确认模式下）、展示识别结果与置信度。',
        '2. 人员管理面板：展示已录入人员列表，支持添加与删除。'
        '添加人员支持多图批量上传，系统自动检测人脸并取平均特征入库。',
        '3. 特征分析模块：选择不同特征提取方式（HOG / 像素统计 / '
        'PCA+DCT+DFT / SVD+LBP / 融合特征），对单张图片实时提取并预览维度与前 10 维数值。',
        '4. 识别记录：显示历史识别记录的时间、姓名、置信度，'
        '支持清空；统计信息实时更新已录入人数与累计识别次数。',
        '5. 系统设置：在前端实时调整识别阈值（0.3 - 0.9），影响后端匹配判定的灵敏度。',
    ]
    # 找到 idx_start 当前的段落元素作为锚点
    anchor_p = doc.paragraphs[idx_start - 1]  # 3.4 接口表后,3.5 标题前
    # 在 idx_start 位置顺序插入
    parent = anchor_p._element.getparent()
    # 重新获取 idx_start 索引(因为可能 anchor_p 后还有别的内容)
    target_p = doc.paragraphs[idx_start]
    target_el = target_p._element
    # 在 target_el 之前依次插入
    # 先把现有 target_p 暂存
    # 把新段落加到文档末尾,然后用 XML 操作挪到前面
    for text in new_paras:
        new_p = doc.add_paragraph(text, style='List Paragraph')
        # 把这个段落挪到 target_el 之前
        new_el = new_p._element
        new_el.getparent().remove(new_el)
        target_el.addprevious(new_el)
    # 删掉旧的占位段
    remove_paragraph(target_p)

# ============================================================
# 5) 修改 4.2.1 代码片段
# ============================================================
# 原 4.2.1 的代码段 P141~P166 替换
idx_421 = find_paragraph_index('FaceRecognitionCore类是系统的核心')
if idx_421 >= 0:
    # 找代码段开始
    idx_code_start = find_paragraph_index('class FaceRecognitionCore:')
    if idx_code_start < 0:
        idx_code_start = idx_421 + 1
    # 找到代码段结尾 (recognize_face 函数结束后)
    # 我们用 4.2.2 标题作为结束边界
    idx_422 = find_paragraph_index('系统实现了完整的图像预处理流程')
    if idx_422 > idx_code_start:
        # 删除从 idx_code_start 到 idx_422-1 的所有段
        for _ in range(idx_422 - idx_code_start):
            p_to_del = doc.paragraphs[idx_code_start]
            remove_paragraph(p_to_del)
        # 插入新代码
        new_code_lines = [
            'class FaceRecognitionCore:',
            '    """人脸识别核心类"""',
            '',
            '    def __init__(self):',
            '        self.detector = dlib.get_frontal_face_detector()',
            '        self.predictor = dlib.shape_predictor(',
            '            \'shape_predictor_68_face_landmarks.dat\')',
            '        self.face_model = dlib.face_recognition_model_v1(',
            '            \'dlib_face_recognition_resnet_model_v1.dat\')',
            '        self.known_faces = []',
            '        self.known_names = []',
            '',
            '    def imread(self, path):',
            '        """读取图片，支持中文路径"""',
            '        data = np.fromfile(path, dtype=np.uint8)',
            '        return cv2.imdecode(data, cv2.IMREAD_COLOR)',
            '',
            '    def detect_faces(self, gray_image):',
            '        """检测人脸（需要灰度图）"""',
            '        return self.detector(gray_image, 1)',
            '',
            '    def extract_features(self, rgb_image, face_rect):',
            '        """提取 128 维人脸特征向量"""',
            '        shape = self.predictor(rgb_image, face_rect)',
            '        descriptor = self.face_model.compute_face_descriptor(',
            '            rgb_image, shape, 1)',
            '        return np.array(descriptor)',
            '',
            '    def recognize_face(self, face_descriptor, th=0.6):',
            '        """辨认（一对多）：返回 (name, similarity, mode)"""',
            '        if not self.known_faces:',
            '            return \'Unknown\', 0.0, \'identification\'',
            '        distances = np.linalg.norm(',
            '            self.known_faces - face_descriptor, axis=1)',
            '        idx = int(np.argmin(distances))',
            '        d = distances[idx]',
            '        sim = 1 - d',
            '        if d < th:',
            '            return self.known_names[idx], sim, \'identification\'',
            '        return \'Unknown\', sim, \'identification\'',
            '',
            '    def verify_face(self, face_descriptor, person_name, th=0.6):',
            '        """确认（一对一）"""',
            '        if person_name not in self.known_names:',
            '            return False, 0.0, \'verification\'',
            '        idx = self.known_names.index(person_name)',
            '        d = float(np.linalg.norm(',
            '            self.known_faces[idx] - face_descriptor))',
            '        return d < th, 1 - d, \'verification\'',
            '',
            '    def add_person_with_preprocess(self, name, image_paths):',
            '        """录入人员：多图取平均特征后 L2 归一化入库"""',
            '        feats = []',
            '        for p in image_paths:',
            '            img = self.imread(p)',
            '            if img is None: continue',
            '            rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)',
            '            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)',
            '            faces = self.detect_faces(gray)',
            '            if not faces: continue',
            '            feats.append(self.extract_features_with_preprocess(',
            '                rgb, faces[0]))',
            '        if not feats: return False',
            '        avg = np.mean(feats, axis=0)',
            '        avg = avg / np.linalg.norm(avg)',
            '        self.known_faces.append(avg)',
            '        self.known_names.append(name)',
            '        return True',
        ]
        # 把代码段插入到 4.2.1 描述段(现在 doc.paragraphs[idx_421])之后
        anchor_el = doc.paragraphs[idx_421]._element
        # 找到 4.2.2 标题段
        idx_422_now = find_paragraph_index('系统实现了完整的图像预处理流程')
        target_el = doc.paragraphs[idx_422_now]._element
        for text in new_code_lines:
            new_p = doc.add_paragraph(text, style='Normal (Web)')
            new_el = new_p._element
            new_el.getparent().remove(new_el)
            target_el.addprevious(new_el)

# ============================================================
# 6) 4.2.2 预处理代码片段 - 改成完整链
# ============================================================
idx_422_des = find_paragraph_index('系统实现了完整的图像预处理流程')
idx_422_code = find_paragraph_index('def preprocess_face(self, image, shape):')
if idx_422_code >= 0:
    # 替换为完整预处理链
    new_422_lines = [
        'def preprocess_face_full(self, img, face_rect):',
        '    """完整预处理：光线补偿→对齐→直方图均衡→归一化→滤波→锐化"""',
        '    rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)',
        '    shape = self.get_face_landmarks(rgb, face_rect)',
        '    img = self.light_compensation(img)',
        '    face_chip = dlib.get_face_chip(img, shape, size=150)',
        '    yuv = cv2.cvtColor(face_chip, cv2.COLOR_BGR2YUV)',
        '    yuv[:, :, 0] = cv2.equalizeHist(yuv[:, :, 0])',
        '    yuv[:, :, 0] = cv2.GaussianBlur(yuv[:, :, 0], (3, 3), 0)',
        '    return cv2.cvtColor(yuv, cv2.COLOR_YUV2RGB), shape',
        '',
        'def preprocess_face(self, img, shape):',
        '    """简化预处理：人脸对齐 + 直方图均衡 + 高斯滤波"""',
        '    face_chip = dlib.get_face_chip(img, shape, size=150)',
        '    yuv = cv2.cvtColor(face_chip, cv2.COLOR_RGB2YUV)',
        '    yuv[:, :, 0] = cv2.equalizeHist(yuv[:, :, 0])',
        '    yuv[:, :, 0] = cv2.GaussianBlur(yuv[:, :, 0], (3, 3), 0)',
        '    return cv2.cvtColor(yuv, cv2.COLOR_YUV2RGB)',
    ]
    # 先把原代码段全部删掉(到 4.2.3 之前)
    idx_423 = find_paragraph_index('人脸录入功能支持从摄像头或图像文件获取人脸图像')
    if idx_423 > idx_422_code:
        for _ in range(idx_423 - idx_422_code):
            p_to_del = doc.paragraphs[idx_422_code]
            remove_paragraph(p_to_del)
        # 在 4.2.3 段前插入
        target_el = doc.paragraphs[idx_423]._element
        for text in new_422_lines:
            new_p = doc.add_paragraph(text, style='Normal (Web)')
            new_el = new_p._element
            new_el.getparent().remove(new_el)
            target_el.addprevious(new_el)

# ============================================================
# 7) 4.2.3 录入功能 - 替换为 add_person_with_preprocess
# ============================================================
idx_423 = find_paragraph_index('人脸录入功能支持从摄像头或图像文件获取人脸图像')
idx_423_code = find_paragraph_index('def add_person(self, name, image_paths):')
if idx_423_code >= 0:
    new_423_lines = [
        'def add_person_with_preprocess(self, name, image_paths):',
        '    """多图平均特征入库（含预处理）"""',
        '    features_list = []',
        '    for img_path in image_paths:',
        '        img = self.imread(img_path)        # 支持中文路径',
        '        if img is None: continue',
        '        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)',
        '        img_gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)',
        '        faces = self.detect_faces(img_gray)',
        '        if not faces: continue',
        '        features = self.extract_features_with_preprocess(',
        '            img_rgb, faces[0])',
        '        features_list.append(features)',
        '    if not features_list: return False',
        '    avg = np.mean(features_list, axis=0)',
        '    avg = avg / np.linalg.norm(avg)        # L2 归一化',
        '    self.known_faces.append(avg)',
        '    self.known_names.append(name)',
        '    return True',
    ]
    idx_424 = find_paragraph_index('实时识别功能从摄像头获取视频流')
    if idx_424 > idx_423_code:
        for _ in range(idx_424 - idx_423_code):
            p_to_del = doc.paragraphs[idx_423_code]
            remove_paragraph(p_to_del)
        target_el = doc.paragraphs[idx_424]._element
        for text in new_423_lines:
            new_p = doc.add_paragraph(text, style='Normal (Web)')
            new_el = new_p._element
            new_el.getparent().remove(new_el)
            target_el.addprevious(new_el)

# ============================================================
# 8) 4.2.4 实时识别 - 改为 Web API 处理流程
# ============================================================
idx_424 = find_paragraph_index('实时识别功能从摄像头获取视频流')
idx_424_code = find_paragraph_index('def process_frame(self, frame):')
if idx_424_code >= 0:
    new_424_lines = [
        '@app.route(\'/api/recognize\', methods=[\'POST\'])',
        'def recognize():',
        '    """Web 接口: 接收上传图片,返回识别结果"""',
        '    mode = request.form.get(\'mode\', \'identification\')',
        '    threshold = float(request.form.get(\'threshold\', 0.6))',
        '    target = request.form.get(\'target\', \'\')',
        '    file = request.files.get(\'image\')',
        '    # 解码图片 + 灰度/RGB 转换',
        '    img_bytes = file.read()',
        '    nparr = np.frombuffer(img_bytes, np.uint8)',
        '    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)',
        '    img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)',
        '    img_gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)',
        '    # 检测人脸并逐一识别',
        '    faces = core.detect_faces(img_gray)',
        '    results = []',
        '    for i, rect in enumerate(faces):',
        '        feats = core.extract_features_with_preprocess(img_rgb, rect)',
        '        if mode == \'verification\':',
        '            ok, sim, _ = core.verify_face(feats, target, threshold)',
            '            results.append({\'name\': target, \'match\': ok,',
            '                            \'similarity\': sim})',
            '        else:',
            '            matches, _ = core.identify_face(feats, threshold)',
            '            best = matches[0] if matches else {\'name\': \'Unknown\',',
            '                                                \'similarity\': 0.0}',
            '            results.append({\'name\': best[\'name\'],',
            '                            \'similarity\': best[\'similarity\']})',
        '    return jsonify({\'success\': True, \'results\': results})',
    ]
    # 删除 4.2.4 整段(到 4.3 之前)
    idx_43 = find_paragraph_index('4.3系统优化')
    if idx_43 < 0:
        idx_43 = find_paragraph_index('系统优化')
    if idx_43 > idx_424_code:
        for _ in range(idx_43 - idx_424_code):
            p_to_del = doc.paragraphs[idx_424_code]
            remove_paragraph(p_to_del)
        # 在 4.3 节标题前插入
        target_el = doc.paragraphs[idx_43]._element
        for text in new_424_lines:
            new_p = doc.add_paragraph(text, style='Normal (Web)')
            new_el = new_p._element
            new_el.getparent().remove(new_el)
            target_el.addprevious(new_el)

# ============================================================
# 9) 5.1 项目成果里把"图形化界面"改成 Web
# ============================================================
idx = find_paragraph_index('设计了简洁、直观的图形化界面')
if idx >= 0:
    replace_paragraph_text(doc.paragraphs[idx],
        '设计了简洁、直观的 Web 可视化界面，支持浏览器端完成录入与识别全流程。')

# ============================================================
# 10) 1.3 任务描述里的"图像采集及检测：可以通过设备进行采集"
# ============================================================
# 不动,本来就是事实

doc.save(DST)
print(f'Saved to: {DST}')
print(f'Size: {os.path.getsize(DST)} bytes')
