"""
补 3 处缺失的示例代码:
1. 4.3.1 末尾: detect_faces 示例 (在 4.3.1 文字之后, 4.3.2 heading 之前)
2. 4.3.2 末尾: 多模板平均示例 (在 4.3.2 文字之后, 4.3.3 文字之前)
3. 5.2 末尾: preprocess + align_face 示例 (在 5.2 文字之后, 5.3 heading 之前)

之前 v2 脚本失败原因: add 末尾 + addprevious 多次时, 第一次创建的就停末尾不动了
这次用"先在末尾 add 2 段, 然后 addprevious 移动"的方式, 验证是否在正确位置
"""
import shutil
import docx
from docx.shared import Pt, RGBColor

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak6'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)


def make_code(text):
    """创建代码段并加入文档, 返回段落对象"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = docx.shared.Pt(14)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    return p


def insert_before(target_p, new_p):
    """把 new_p 移动到 target_p 之前"""
    target_p._element.addprevious(new_p._element)


# ============================================================
# 1. 4.3.1 末尾补 detect_faces
# ============================================================
h_432 = None
for p in paras:
    if p.text.strip() == '4.3.2识别准确率优化' and p.style.name.startswith('Heading'):
        h_432 = p
        break
if h_432:
    detect_code = (
        "def detect_faces(self, gray_image, upsample=1):\n"
        "    return self.detector(gray_image, upsample)"
    )
    new_p = make_code(detect_code)
    insert_before(h_432, new_p)
    print('1. 4.3.1 末尾补 detect_faces ✓')


# ============================================================
# 2. 4.3.2 末尾补多模板平均
# ============================================================
h_433 = None
for p in paras:
    if p.text.strip() == '4.3.3中文支持优化' and p.style.name.startswith('Heading'):
        h_433 = p
        break
if h_433:
    multi_code = (
        "def add_person_with_preprocess(self, name, image_paths):\n"
        "    feats = []\n"
        "    for p in image_paths:\n"
        "        img = self.imread(p)\n"
        "        if img is None: continue\n"
        "        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)\n"
        "        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)\n"
        "        faces = self.detect_faces(gray)\n"
        "        if not faces: continue\n"
        "        feats.append(self.extract_features_with_preprocess(rgb, faces[0]))\n"
        "    if not feats: return False\n"
        "    avg = np.mean(feats, axis=0)\n"
        "    avg = avg / np.linalg.norm(avg)\n"
        "    self.known_faces.append(avg)\n"
        "    self.known_names.append(name)\n"
        "    return True"
    )
    new_p = make_code(multi_code)
    insert_before(h_433, new_p)
    print('2. 4.3.2 末尾补 add_person_with_preprocess ✓')


# ============================================================
# 3. 5.2 末尾补 preprocess + align_face
# ============================================================
h_53 = None
for p in paras:
    if p.text.strip() == '5.3项目难点与解决方案' and p.style.name.startswith('Heading'):
        h_53 = p
        break
if h_53:
    pre_code = (
        "def preprocess_pipeline(self, image):\n"
        "    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
        "    gray = cv2.equalizeHist(gray)\n"
        "    gray = cv2.GaussianBlur(gray, (5, 5), 0)\n"
        "    return gray"
    )
    align_code = (
        "def align_face(self, rgb_image, face_rect):\n"
        "    shape = self.predictor(rgb_image, face_rect)\n"
        "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
        "    return chip"
    )
    # 先创建两段, 都 add 到末尾
    p_pre = make_code(pre_code)
    p_align = make_code(align_code)
    # 重要: 先移动 p_align 到 h_53 前 (它在 p_pre 之后, 文档末尾倒数第二)
    insert_before(h_53, p_align)
    # 再移动 p_pre 到 h_53 前 (此时 h_53 前已经有 p_align)
    insert_before(h_53, p_pre)
    # 最终顺序: ...其他段... p_pre, p_align, h_53
    print('3. 5.2 末尾补 preprocess + align_face ✓')


d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')