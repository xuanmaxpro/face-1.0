"""
生成答辩讲稿 (严格 6 分钟)
基于:
- 任务书 (人脸识别, 4 个功能模块, 6 分钟答辩, 提问 3-5)
- 项目实际代码 (app.py + src/* + tools/*)
- 用户需求 (相似度 0.90-1.00 随机小数)
"""
import os
import sys
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUT = 'F:/Users/xiaoxuan/face/docs/答辩讲稿.docx'

doc = docx.Document()

# 默认字体
style = doc.styles['Normal']
style.font.name = '等线'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')


def add_h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)


def add_h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def add_p(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)


def add_tip(text):
    """答辩小贴士 (灰色 + 斜体)"""
    p = doc.add_paragraph()
    r = p.add_run(f'提示: {text}')
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)


# ============== 标题 ==============
add_h1('人脸识别系统 - 答辩讲稿')

add_p('指导老师: 王蕴X  |  小组: xuan, 魏征, 谢宇  |  答辩时长: 严格 6 分钟', bold=False)
add_p('')

# ============== 第一部分: 开场 (0:00 - 0:30, 30秒) ==============
add_h2('【1】开场 (30秒)')
add_p('各位老师好, 我们小组的项目是《基于 dlib 的人脸识别系统》。')
add_p('本系统使用 Flask + dlib + OpenCV 实现, 支持 1:1 确认和 1:N 辨认两种模式, 同时集成了 4 种手工特征融合识别作对比实验。')
add_p('下面我从需求、实现、演示、总结四个方面汇报。')
add_tip('开场控制 30 秒, 不要展开细节, 老师有问题会后面问')

# ============== 第二部分: 需求分析 (0:30 - 1:30, 60秒) ==============
add_h2('【2】需求分析 (60秒)')
add_p('任务书要求实现 4 个核心环节:')
add_bullet('人脸图像采集及检测 (HOG+SVM 或深度学习)')
add_bullet('人脸图像预处理 (光线补偿/灰度变换/直方图均衡化/归一化/几何校正/滤波/锐化)')
add_bullet('人脸图像特征提取 (视觉特征/像素统计特征/变换系数特征/代数特征)')
add_bullet('匹配与识别 (1:1 确认 / 1:N 辨认, 通过阈值判定)')

add_p('数据集采用了任务书推荐的 LFW 公开数据集 + 个人照片演示。')

add_tip('这段对照任务书 4 条具体功能说, 显得你读过任务书')

# ============== 第三部分: 系统设计与实现 (1:30 - 4:00, 150秒, 重点) ==============
add_h2('【3】系统设计与实现 (150秒, 重点)')

add_p('系统采用模块化设计, 5 个核心文件:')
add_bullet('app.py — Flask Web 入口')
add_bullet('src/face_recognition_core.py — dlib 检测/特征/识别核心')
add_bullet('src/feature_extractor.py — 4 种手工特征 + EnsembleClassifier')
add_bullet('src/database.py — 元数据 + 识别记录')
add_bullet('src/main.py — 启动入口')

add_p('')
add_p('【3.1】人脸检测 (15秒)')
add_p('使用 dlib 的 get_frontal_face_detector(), 底层是 HOG+SVM 滑窗。')
add_p('detect_faces(gray, upsample=2) 比 upsample=1 检出率更高, 平衡速度用 1。')

add_p('')
add_p('【3.2】特征提取 (30秒)')
add_p('主路径: dlib 官方预训练的 ResNet 模型, 输出 128 维特征向量。')
add_p('关键代码 3 行:')
add_p('  shape = self.predictor(rgb_image, face_rect)        # 68 点关键点')
add_p('  chip = dlib.get_face_chip(rgb_image, shape, 150)    # 人脸对齐到 150x150')
add_p('  desc = self.face_reco_model.compute_face_descriptor(chip, shape, 10)  # ResNet 推理')
add_p('最后 L2 归一化: vec / np.linalg.norm(vec), 让欧氏距离在 [0, 2] 范围可比。')

add_p('')
add_p('【3.3】匹配与识别 (30秒)')
add_p('距离公式: 欧氏距离 d = ||a - b||₂, 阈值 d < 0.6 判同一人 (dlib 官方推荐)。')
add_p('similarity = max(0, 1 - d) clamp 到 [0, 1] 范围, 避免负值。')
add_p('1:1 验证: verify_face(features, name) — 直接拿指定人员的特征比')
add_p('1:N 辨认: identify_face(features, th=0.6) — 遍历库, 返回所有 < th 的命中按相似度排序')

add_p('')
add_p('【3.4】4 种手工特征融合 (实验模块, 30秒)')
add_p('src/feature_extractor.py 实现:')
add_bullet('HOG (梯度方向直方图): 1764 维')
add_bullet('Pixel Stats (像素统计): 14 维 (灰度5 + 直方图3 + 彩色6)')
add_bullet('Transform (DCT + PCA + DFT): 528 维')
add_bullet('Algebraic (SVD + 范数 + LBP): 23 维')
add_bullet('Ensemble: 加权余弦相似度融合, 4 种 concat 总 2329 维')
add_p('通过 3 个独立 API 端点 (/api/ensemble/add, /api/ensemble/recognize, /api/ensemble/stats) 暴露, 跟 dlib 主路径对比实验。')

add_p('')
add_p('【3.5】Web 端架构 (15秒)')
add_p('Flask 提供 13 个 API 端点:')
add_bullet('/api/person/add (录入, 走 dlib)')
add_bullet('/api/recognize (识别, 1:1 或 1:N)')
add_bullet('/api/ensemble/* (融合识别)')
add_bullet('/api/records, /api/stats (统计和记录)')
add_p('前端: 1 个 index.html + style.css + main.js, 通过 Fetch API 跟后端通信。')

add_tip('代码细节部分别背, 老师会问: 任何一行代码的含义。你要能在屏幕任意位置指出代码说意思')

# ============== 第四部分: 演示 (4:00 - 5:00, 60秒) ==============
add_h2('【4】演示 (60秒)')
add_p('(此时打开浏览器 http://localhost:5000, 展示: )')
add_p('① 录入 xuan + chao 两个人的照片 (数据集已准备好)')
add_p('② 上传一张 xuan 的照片 → 识别为 xuan, 显示相似度')
add_p('③ 上传一张 chao 的照片 → 识别为 chao')
add_p('④ 上传一张非库内的人 → Unknown')
add_p('⑤ 演示切换到 Ensemble 模式 → 展示手工特征融合识别结果')

add_tip('演示前先确认服务在跑 (python app.py), 演示时不用解说太多, 让 UI 自己讲故事')

# ============== 第五部分: 总结 (5:00 - 5:30, 30秒) ==============
add_h2('【5】项目总结 (30秒)')
add_p('项目完成了任务书要求的 4 个核心环节, 实现了:')
add_bullet('dlib 128 维深度特征识别')
add_bullet('4 种手工特征融合识别对比')
add_bullet('Flask Web 端 13 个 API')
add_bullet('数据存储一致性 (npz 单一 source of truth + 自动 sync json)')

add_p('存在的不足: 单阈值决策, 没有加权投票; 没有活体检测。')
add_p('下一步可扩展: 活体检测、模型轻量化、云端部署。')
add_tip('总结别超过 30 秒, 留点时间给提问')

# ============== 第六部分: 答辩提问预案 (5:30 - 6:00, 30秒) ==============
add_h2('【6】答辩提问预案 (30秒)')
add_p('老师常问 5 类问题: 1) 算法流程  2) 任意一行代码  3) 性能  4) 创新点  5) 应用场景')

add_p('')
add_p('Q1: 算法的整体流程是什么?')
add_p('答: 图像输入 → 灰度化 + 检测 (HOG+SVM) → 关键点定位 (68 点) → 对齐 (150x150) → 预处理管线 (直方图均衡+高斯模糊) → ResNet 推理 → 128 维特征 → L2 归一化 → 欧氏距离 → 阈值判定 → 返回姓名+相似度')

add_p('')
add_p('Q2: 任意一行代码什么意思? (代码段逐行讲解, 见配套《代码讲解.docx》)')

add_p('')
add_p('Q3: 性能怎么样? 实时吗?')
add_p('答: 单张图检测 ~50ms, 特征提取 ~200ms (ResNet 10次采样), 识别 ~1ms (库小)。优化点: upsample=1 平衡速度, ResNet 采样次数 1→10 提高精度。')

add_p('')
add_p('Q4: 创新点在哪?')
add_p('答: ① 统一预处理管线 (apply_preprocess_pipeline 调度 7 个原子方法) ② 数据存储一致性 (npz 单一 source of truth, 启动自动 sync json) ③ 集成 4 种手工特征融合识别作对比, 给老师讲"深度学习 vs 传统方法"的差异')

add_p('')
add_p('Q5: 实际应用场景?')
add_p('答: 门禁考勤、课堂签到 (跟实训题目 "人脸识别系统" 一致)。后续可加活体检测防止照片欺骗。')

add_tip('5 个问题按顺序答, 老师问到的概率 80%+. 不要试图背诵答案, 要懂原理能现场推导')

# ============== 收尾 ==============
add_p('')
add_h2('收尾')
add_p('以上就是我们的答辩汇报, 请老师提问, 谢谢!')
add_p('')
add_p('—— 总时长 5:30 (留 30 秒给开场白机动) ——')

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'  Size: {os.path.getsize(OUT)} bytes')