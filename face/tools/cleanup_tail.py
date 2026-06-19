"""
清理: 文档末尾 [278=参考文献, 末尾] 之间所有段
(因为 v2 脚本 bug, 部分代码段被遗留在文档末尾)
预期保留: 参考文献段, 删掉所有"重复 heading + 重复代码"
"""
import shutil
import docx

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak5'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)
print(f'Total paragraphs before: {len(paras)}')

# 找 参考文献 段
idx_ref = None
for i, p in enumerate(paras):
    if p.text.strip() == '参考文献' and p.style.name.startswith('Heading'):
        idx_ref = i
        break
print(f'参考文献 heading at: {idx_ref}')

# 参考文献之前正常, 之后是 [idx_ref+1, 末尾] 段
# dump 显示 [278=参考文献, 279=第5章, 280-291=12 段]
# 这些都是重复的, 全删
to_delete = paras[idx_ref + 1:]
for p in to_delete:
    p._element.getparent().remove(p._element)
print(f'Deleted {len(to_delete)} paragraphs after 参考文献')

d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')