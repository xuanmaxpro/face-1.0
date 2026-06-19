"""
最后补 5.2 末尾的代码 + 重建 5.3 heading (在 5.3.1 之前)
"""
import shutil
import docx
from docx.shared import Pt, RGBColor

SRC = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx'
BAK = 'F:/Users/xiaoxuan/face/docs/实训报告_修改版.docx.bak8'

shutil.copy2(SRC, BAK)
print(f'Backup: {BAK}')

d = docx.Document(SRC)
paras = list(d.paragraphs)


def make_heading(text, level=2):
    p = d.add_paragraph()
    p.style = d.styles[f'Heading {level}']
    p.add_run(text)
    return p


def make_code(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = docx.shared.Pt(14)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(50, 50, 50)
    return p


def insert_before(target_p, new_p):
    target_p._element.addprevious(new_p._element)


# 重建 5.3 heading (Heading 2) 在 5.3.1 之前
h_531 = None
for p in d.paragraphs:
    if p.text.strip() == '5.3.1识别准确率问题' and p.style.name == 'Heading 3':
        h_531 = p
        break

if h_531:
    h_53 = make_heading('5.3项目难点与解决方案', level=2)
    insert_before(h_531, h_53)
    print('1. 重建 5.3 heading (Heading 2) ✓')
else:
    print('1. ✗ 找不到 5.3.1 heading')

# 5.2 末尾补 preprocess + align_face (在 5.3 heading 前)
if h_531:
    pre_code = (
        "def preprocess_pipeline(self, image):\n"
        "    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)\n"
        "    gray = cv2.equalizeHist(gray)\n"
        "    gray = cv2.GaussianBlur(gray, (5, 5), 0)\n"
        "    return gray"
    )
    align_code = (
        "def align_face(self, rgb_image, face_rect):\n"
        "    shape = self.predictor(rgb_image, face_rect)\n"
        "    chip = dlib.get_face_chip(rgb_image, shape, size=150)\n"
        "    return chip"
    )
    p_pre = make_code(pre_code)
    p_align = make_code(align_code)
    insert_before(h_53, p_align)
    insert_before(h_53, p_pre)
    print('2. 5.2 末尾补 preprocess + align_face ✓')

d.save(SRC)

import os
print(f'\nSaved: {SRC}')
print(f'  Size: {os.path.getsize(SRC)} bytes')
print(f'  Total paragraphs: {len(docx.Document(SRC).paragraphs)}')